from PyPDF2 import PdfReader
from docx import Document
from tkinter import Tk, filedialog

def pdf_to_word():
    # Hide the root tkinter window
    Tk().withdraw()

    # Ask user to select the PDF file
    pdf_path = filedialog.askopenfilename(
        title="Select PDF file",
        filetypes=[("PDF Files", "*.pdf")]
    )
    if not pdf_path:
        print("No file selected. Exiting...")
        return

    # Ask for output Word file path
    doc_path = filedialog.asksaveasfilename(
        title="Save Word file as",
        defaultextension=".doc",
        filetypes=[("Word Files", "*.doc"), ("All Files", "*.*")]
    )
    if not doc_path:
        print("No output path selected. Exiting...")
        return

    # Read the PDF
    reader = PdfReader(pdf_path)
    document = Document()

    print("Extracting text from PDF...")

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text:
            # Preserve punctuation and structure
            document.add_paragraph(text.strip())
        else:
            document.add_paragraph(f"[No extractable text on page {page_num}]")

    # Save to .doc file
    document.save(doc_path)
    print(f"✅ Successfully extracted and saved to: {doc_path}")

if __name__ == "__main__":
    pdf_to_word()
